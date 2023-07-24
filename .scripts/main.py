import os
import markdown
from docx import Document

def md_to_docx(md_file_path, docx_file_path):
    # Create a new Document
    doc = Document()
    
    # Read the markdown file
    with open(md_file_path, 'r') as md_file:
        md_content = md_file.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)

    # Add the HTML content to the document
    doc.add_paragraph(html_content)
    
    # Save the document
    doc.save(docx_file_path)

def merge_md_files_to_docx(md_dir_path, docx_file_path):
    # Get all .md files in the directory
    md_files = [f for f in os.listdir(md_dir_path) if f.endswith('.md')]
    
    # Create a new Document
    doc = Document()

    # Loop through all .md files
    for md_file in md_files:
        # Read the markdown file
        with open(os.path.join(md_dir_path, md_file), 'r') as f:
            md_content = f.read()

        # Convert markdown to HTML
        html_content = markdown.markdown(md_content)

        # Add the HTML content to the document
        doc.add_paragraph(html_content)
        # Add a page break after each file
        doc.add_page_break()

    # Save the document
    doc.save(docx_file_path)

# Use the function
merge_md_files_to_docx('../A-Song-of-Ice-and-Fire/The-World-of-Ice-and-Fire', '../A-Song-of-Ice-and-Fire/The-World-of-Ice-and-Fire/output.docx')

""" from markdown2docx import Markdown2Docx

def merge_md_files_to_docx(md_dir_path, docx_file_path):
    # Get all .md files in the directory
    md_files = [f for f in os.listdir(md_dir_path) if f.endswith('.md')]
    
    # Initialize a Markdown2Docx converter
    converter = Markdown2Docx()

    # Loop through all .md files
    for md_file in md_files:
        # Convert and append markdown file to the docx file
        converter.convert(os.path.join(md_dir_path, md_file), docx_file_path, append=True)

# Use the function
merge_md_files_to_docx('../A-Song-of-Ice-and-Fire/The-World-of-Ice-and-Fire', '../A-Song-of-Ice-and-Fire/The-World-of-Ice-and-Fire/output.docx') """
